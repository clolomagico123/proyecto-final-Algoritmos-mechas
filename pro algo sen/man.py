import os
import openpyxl
import mimetypes
import smtplib
import ssl
from email.message import EmailMessage
from docx import Document



directorio = "C:\\Users\\carlos\\Documents\\parcial"
archivo = open(directorio + "\\preciosdelocura.txt", "w")
archivo.close()




#Con un menu asignamos las funciones de nuestro codigo 
def menu():
    print('============================')
    print('Precios de Locura')
    print('============================')
    print('Productos')
    print('1 - Agregar producto')
    print('2 - Editar producto')
    print('3 - Eliminar producto')
    print('4 - listar productos')
    print('5 - Enviar cotización por correo')
    print('============================')
    print('Productos')
    print('6 - Agregar cliente')
    print('7 - Editar cliente')
    print('8 - Eliminar cliente')
    print('9 - Listar clientes')
    print('============================')
    print('Pedidos')
    print('10 - Agregar pedido')
    print('11 - Eliminar pedido')
    print('12- Listar pedidos')
    print('============================')
    print('Informes')
    print('13 - Total de venta por cliente')
    print('14 - Total de ventas por producto')
    print('============================')
    print('Varios')
    print('15 - Crear copia de seguridad de datos')
    print('0 - Salir')
    print('============================')
    option = input('Introduzca el número de la opción deseada: ')
    return option
#esta es la otra parte del menu aca es el codigo
def directory():
    while True:
        option = menu()
       
        if option == '1':
            nom1 = input('Introduce el nombre del Producto: ')
            prec2 = input('Introduce el precio del producto: ')
            exis3 = input('Introduce la existencia del producto: ')
           
            libro = openpyxl.load_workbook('inventario.xlsx')
            # crear una nueva hoja
            hoja = libro['productos']

            produinv= [
                 {
             nom1,
             prec2,
             exis3,
                },
                ]
            primeraFila = hoja.max_row + 1
            for venta in produinv:
                hoja['A' + str(primeraFila)].value = nom1
                hoja['B' + str(primeraFila)].value = prec2
                hoja['C' + str(primeraFila)].value = exis3
                primeraFila = primeraFila + 1
            #guardar cambios
            libro.save("inventario.xlsx")
            
        
        elif option == '2':
                
                Editarinp = int(input("Ingrese el numero de linea a editar:  "))
                libro = openpyxl.load_workbook('inventario.xlsx')
    
                hoja = libro['productos']
                Edita = Editarinp + 1
                hoja['A1'].value = "Producto"
                hoja['B1'].value = "Precio"
                hoja['C1'].value = "Existencia"
                agregarEditar = []
                Agrediccionarioedit = {}
                ProductoEdit = input("nuevo nombre del producto:  "+ hoja['A' + str(Edita)].value +"\n")
                Agrediccionarioedit['producto'] = ProductoEdit
                PrecioEdit = float(input("nuevo precio del producto:  "+ str(hoja['B' + str(Edita)].value) +"\n"))
                Agrediccionarioedit['precio'] = PrecioEdit
                ExEdit = int(input("nuevo cantidad existente del producto:  "+ str(hoja['C' + str(Edita)].value) +"\n"))
                Agrediccionarioedit['existencia'] = ExEdit
                agregarEditar.append(Agrediccionarioedit)
                for producEdit in agregarEditar:
                    hoja['A' + str(Edita)].value = producEdit['producto']
                    hoja['B' + str(Edita)].value = producEdit['precio']
                    hoja['C' + str(Edita)].value = producEdit['existencia']
                    libro.save("inventario.xlsx")
        
        elif option == '3':
                eliminarSelect = int(input("Ingrese el numero de linea a eliminar:  "))
                libro = openpyxl.load_workbook('inventario.xlsx')
                
                hoja = libro['productos']
                elim = eliminarSelect + 1
               
                hoja['A1'].value = "Producto"
                print("linea eliminada del producto:  "+ hoja['A' + str(elim)].value +"\n")
                hoja.delete_rows(idx=elim,amount=3)
                
                libro.save("inventario.xlsx")
           
        elif option == '4':

            libro = openpyxl.load_workbook('inventario.xlsx')
            hoja = libro['productos']

            hojaClientes = libro['productos']

            diccionarioCliente = {}
            clientes = []
            for row in range(2, hojaClientes.max_row + 1):
                # explora fila por fila
                diccionarioCliente['nombre'] = hojaClientes["A" + str(row)].value
                diccionarioCliente['precio'] = hojaClientes["B" + str(row)].value
                diccionarioCliente['existencia'] = hojaClientes["C" + str(row)].value
                clientes.append(diccionarioCliente)
                diccionarioCliente = {}

            for item in clientes:
                print ("Nombre: " + item['nombre'])
                print ("precio: " + item['precio'])
                print ("existencia: " + str(item['existencia']))
                
            

        elif option == '5':
            nom88 = input('Introduce tu nombre: ')
            corre1 = input('Introduce tu correo: ')
            proc99 = input('producto a cotizar: ')
            
            document = Document()
            document.add_heading('Cotizar', 0)
            p = document.add_paragraph('Estimado cliente '+nom88)
            document.add_paragraph('el Precio de nuestro producto '+proc99)
            document.add_paragraph('es: 5')
            document.save('cotización.docx')

            DIRECCION_DEL_SERVIDOR = "smtp.gmail.com"
            PUERTO = 587
            DIRECCION_DE_ORIGEN = "marrana091@gmail.com"
            CONTRASENA = "pastel123"

            #Contenido del mensaje
            mensaje = EmailMessage()
            mensaje["Subject"] = "Copia de seguridad Proyecto"
            mensaje["From"] = DIRECCION_DE_ORIGEN
            mensaje["To"] = corre1

            mensaje.set_content("hola solicitastes una cotización")

            mensaje.add_alternative(""" 
            <p> 
                <h1>cotización</h1>
                este es el documento de word de la cotización
            </p>
            """, subtype = "html")

            nombre_de_archivo = "cotización.docx"
            ctype, encoding = mimetypes.guess_type(nombre_de_archivo)

            if ctype is None or encoding is not None:
                ctype = 'application/octet-stream'

            tipoPrincipal, subTipo = ctype.split('/', 1)

            with open(nombre_de_archivo, 'rb') as archivoLeido:
                mensaje.add_attachment(archivoLeido.read(), maintype=tipoPrincipal, subtype = subTipo, filename = nombre_de_archivo)

           
            smtp = smtplib.SMTP(DIRECCION_DEL_SERVIDOR, PUERTO)
            smtp.starttls()
            smtp.login(DIRECCION_DE_ORIGEN, CONTRASENA)
            smtp.send_message(mensaje)
            #la tupla es un tipo de datos que permite devolver dos o mas valores de una funcion
            #def mifunction():
                #return 0, 100 







        elif option == '6':
            nom3 = input('Introduce tu nombre: ')
            nit = input('Introduce tu nit: ')
            direc = input('introduce tu direccion ')
            
            libro = openpyxl.load_workbook('inventario.xlsx')
            # crear una nueva hoja
            hoja = libro['clientes']

            produinv= [
                 {
             nom3,
             nit,
             direc,
                },
                ]
            primeraFila = hoja.max_row + 1
            for venta in produinv:
                hoja['A' + str(primeraFila)].value = nom3
                hoja['B' + str(primeraFila)].value = nit
                hoja['C' + str(primeraFila)].value = direc
                primeraFila = primeraFila + 1
            #guardar cambios
            libro.save("inventario.xlsx")

        elif option == '7':
                Editarinp = int(input("Ingrese el numero de linea a editar:  "))
                libro = openpyxl.load_workbook('inventario.xlsx')
    
                hoja = libro['clientes']
                Edita = Editarinp + 1
                hoja['A1'].value = "Nombre"
                hoja['B1'].value = "Nit"
                hoja['C1'].value = "Direccion"
                agregarEditar = []
                Agrediccionarioedit = {}
                ProductoEdit = input("nuevo nombre del Cliente:  "+ hoja['A' + str(Edita)].value +"\n")
                Agrediccionarioedit['producto'] = ProductoEdit
                PrecioEdit = float(input("nuevo Nit del cliente:  "+ str(hoja['B' + str(Edita)].value) +"\n"))
                Agrediccionarioedit['precio'] = PrecioEdit
                ExEdit = input("nueva direccion del cliente:  "+ str(hoja['C' + str(Edita)].value) +"\n")
                Agrediccionarioedit['existencia'] = ExEdit
                agregarEditar.append(Agrediccionarioedit)
                for producEdit in agregarEditar:
                    hoja['A' + str(Edita)].value = producEdit['producto']
                    hoja['B' + str(Edita)].value = producEdit['precio']
                    hoja['C' + str(Edita)].value = producEdit['existencia']
                    libro.save("inventario.xlsx")
    
        elif option == '8':
                eliminarSelect = int(input("Ingrese el numero de linea a eliminar del cliente:  "))
                libro = openpyxl.load_workbook('inventario.xlsx')
                
                hoja = libro['clientes']
                elim = eliminarSelect + 1
               
                hoja['A1'].value = "Nombre"
                print("linea eliminada del cliente:  "+ hoja['A' + str(elim)].value +"\n")
                hoja.delete_rows(idx=elim,amount=3)
                
                libro.save("inventario.xlsx")

        elif option == '9':
            libro = openpyxl.load_workbook('inventario.xlsx')
            hoja = libro['clientes']

            hojaClientes = libro['clientes']

            diccionarioCliente = {}
            clientes = []
            for row in range(2, hojaClientes.max_row + 1):
                # explora fila por fila
                diccionarioCliente['nombre'] = hojaClientes["A" + str(row)].value
                diccionarioCliente['nit'] = hojaClientes["B" + str(row)].value
                diccionarioCliente['direccion'] = hojaClientes["C" + str(row)].value
                clientes.append(diccionarioCliente)
                diccionarioCliente = {}

            for item in clientes:
                print ("Nombre: " + item['nombre'])
                print ("nit: " + item['nit'])
                print ("direccion: " + str(item['direccion']))

        elif option == '10':
            nom5 = input('Introduce tu nombre: ')
            prope = input('Introduce tu producto: ')
            canpro = input('Introduce cantidad del producto: ')
            val = input('valor del pedido: ')

            libro = openpyxl.load_workbook('inventario.xlsx')
            # crear una nueva hoja
            hoja = libro['pedidos']

            produinv= [
                 {
             nom5,
             prope,
             canpro,
            val,
                },
                ]
            primeraFila = hoja.max_row + 1
            for venta in produinv:
                hoja['A' + str(primeraFila)].value = nom5
                hoja['B' + str(primeraFila)].value = prope
                hoja['C' + str(primeraFila)].value = canpro
                hoja['D' + str(primeraFila)].value = val
                primeraFila = primeraFila + 1
            #guardar cambios
            libro.save("inventario.xlsx")


        elif option == '11':
                eliminarSelect = int(input("Ingrese el numero de linea para eliminar el pedido:  "))
                libro = openpyxl.load_workbook('inventario.xlsx')
                
                hoja = libro['pedidos']
                elim = eliminarSelect + 1
               
                hoja['A1'].value = "nom cliente"
                print("linea eliminada del cliente:  "+ hoja['A' + str(elim)].value +"\n")
                hoja.delete_rows(idx=elim,amount=4)
                
                libro.save("inventario.xlsx")  

        elif option == '12':
            libro = openpyxl.load_workbook('inventario.xlsx')
            hoja = libro['pedidos']

            hojaClientes = libro['pedidos']

            diccionarioCliente = {}
            clientes = []
            for row in range(2, hojaClientes.max_row + 1):
                # explora fila por fila
                diccionarioCliente['nom cliente'] = hojaClientes["A" + str(row)].value
                diccionarioCliente['nom pro'] = hojaClientes["B" + str(row)].value
                diccionarioCliente['cantidad pro'] = hojaClientes["C" + str(row)].value
                diccionarioCliente['valor pedido'] = hojaClientes["D" + str(row)].value
                clientes.append(diccionarioCliente)
                diccionarioCliente = {}

            for item in clientes:
                print ("Nombre: " + item['nom cliente'])
                print ("producto: " + item['nom pro'])
                print ("cantidad producto: " + str(item['cantidad pro']))
                print ("valor del pedido: " + str(item['valor pedido']))            
        
        elif option == '13':
            nom = input('Introduce el nombre del cliente: ') 

        elif option == '14':
            nom = input('Introduce el nombre del producto: ') 

        elif option == '15':
            
            DIRECCION_DEL_SERVIDOR = "smtp.gmail.com"
            PUERTO = 587
            DIRECCION_DE_ORIGEN = "marrana091@gmail.com"
            CONTRASENA = "pastel123"

            #Contenido del mensaje
            mensaje = EmailMessage()
            mensaje["Subject"] = "Copia de seguridad Proyecto"
            mensaje["From"] = DIRECCION_DE_ORIGEN
            mensaje["To"] = "carlosalfonso2602@gmail.com"

            mensaje.set_content("Este es el cuerpo del mensaje")

            mensaje.add_alternative(""" 
            <p> 
                <h1>Copia de Seguridad</h1>
                este es el documento de excel que respalda todo el inventario 
            </p>
            """, subtype = "html")

            nombre_de_archivo = "inventario.xlsx"
            ctype, encoding = mimetypes.guess_type(nombre_de_archivo)

            if ctype is None or encoding is not None:
                ctype = 'application/octet-stream'

            tipoPrincipal, subTipo = ctype.split('/', 1)

            with open(nombre_de_archivo, 'rb') as archivoLeido:
                mensaje.add_attachment(archivoLeido.read(), maintype=tipoPrincipal, subtype = subTipo, filename = nombre_de_archivo)

           
            smtp = smtplib.SMTP(DIRECCION_DEL_SERVIDOR, PUERTO)
            smtp.starttls()
            smtp.login(DIRECCION_DE_ORIGEN, CONTRASENA)
            smtp.send_message(mensaje)
            #la tupla es un tipo de datos que permite devolver dos o mas valores de una funcion
            #def mifunction():
                #return 0, 100       
                    


        #y esta seria la ultima opcion que saldriamos del ciclo con un break
        else:
            break
    return
directory()